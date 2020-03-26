from selenium import webdriver
from selenium.webdriver.common.keys import Keys
import time
import xlrd
import xlsxwriter
import os
import pyscreenshot as ImageGrab
import pyautogui as py
import pynput
from pynput.keyboard import Key, Controller

#TWO PRESS DOWN

path="C:/Users/SHBG7410/Desktop/DILIP/Rishika/gore_main_excel.xlsx"

Wb=xlrd.open_workbook(path)

p9="C:/Users/SHBG7410/Desktop/DILIP/Rishika/TEST/Images/Capture.png"

ss=Wb.sheet_by_index(0)
var=ss.nrows
var2=ss.ncols
print(var,"",var2)
global i,kk
time.sleep(2)

#DOWN ARROW
#for i in range(1,5):
 #   py.click(1357,718)

for x in range (var2):
    dev=ss.col_values(x)
    kk=[]
    for nk in range(1,var):
        print("nk",nk)
        kk=dev[nk]
        print("kk",kk)
        py.click(89,439)
        time.sleep(1)
        py.typewrite(kk)
        #print("device",kk)
#SEARCH
        time.sleep(1)
        py.click(338,438)
        time.sleep(1)
#VIEW
        
        py.click(1207,490)
        time.sleep(30)
#Month
        py.click(1350,91)
        time.sleep(1)
#Previous
        py.click(1149,262)
        #py.click(216,221)
        time.sleep(25)
        

#LAST MONTH        
        
        #py.click(991,88)
        im=ImageGrab.grab(bbox=(190,70,1360,720))
        #im.show()
        d=str(nk)
        im.save(d+".png")
        #print(d+".png")

        py.click(430,16)
        
        time.sleep(1)
        try:
            ca=py.locateOnScreen(p9)
            ca=py.center(ca)
            ca1,ca2=ca
            py.click(ca1,ca2)
        except:
            continue        

#print("done")

#++++++++++Gecko shutdown_++++++
#os.system('tskill plugin-container')
#driver.quit()

#++++++++++++++ PASTE IN DOC +++++++++++++++

from docx import Document
from docx.shared import Inches
d=Document()

d.add_heading('GORE IP Bandwidth Utilization GRAPHS',0)
#d.add_paragraph('JE Americas',style='List Bullet')
d.add_paragraph('Monthly- GRAPHS',style='List Bullet')


for r in range(0,1):
    dev2=ss.col_values(r)
    kk2=[]
    for nk2 in range(1,var):
        st=str(nk2)
        kk2=dev2[nk2]
        #print(kk2)
        d.add_paragraph(kk2,style='List Bullet')
        d.add_picture(st+".png",width=Inches(6.77))


d.save("Gore Bandwidth Utilization.docx")
print("Documrnt Created")

